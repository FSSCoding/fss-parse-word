#!/usr/bin/env python3
"""
Word - Global Document Conversion Tool
Safe, hash-validated bidirectional conversion between .docx and .md formats.

Author: Isabella (Testing & Validation Specialist)
Safety Features: Hash checking, collision detection, confirmation prompts
"""

import argparse
import hashlib
import json
import re
import sys
import yaml
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict
import tempfile
import shutil

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.dml import MSO_THEME_COLOR_INDEX
    from docx.oxml.ns import qn
except ImportError:
    print("❌ Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import markdown
    from markdown.extensions import codehilite, tables, toc, fenced_code
except ImportError:
    print("❌ Error: markdown not installed. Run: pip install markdown")
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("⚠️  Warning: PyYAML not installed. YAML config support disabled. Install with: pip install PyYAML")
    yaml = None


@dataclass
class SafetyConfig:
    """Configuration for safety mechanisms."""
    require_confirmation: bool = True
    create_backup: bool = True
    check_hash: bool = True
    prevent_overwrite: bool = True
    backup_suffix: str = ".backup"


@dataclass
class FormatMetadata:
    """Stores formatting information for reconstruction."""
    bold_ranges: List[Tuple[int, int]] = None
    italic_ranges: List[Tuple[int, int]] = None
    heading_levels: Dict[int, int] = None
    lists: List[Dict[str, Any]] = None
    tables: List[Dict[str, Any]] = None
    hyperlinks: List[Dict[str, Any]] = None
    images: List[Dict[str, Any]] = None
    styles: Dict[str, str] = None
    file_hash: str = ""
    conversion_timestamp: str = ""
    
    def __post_init__(self):
        if self.bold_ranges is None:
            self.bold_ranges = []
        if self.italic_ranges is None:
            self.italic_ranges = []
        if self.heading_levels is None:
            self.heading_levels = {}
        if self.lists is None:
            self.lists = []
        if self.tables is None:
            self.tables = []
        if self.hyperlinks is None:
            self.hyperlinks = []
        if self.images is None:
            self.images = []
        if self.styles is None:
            self.styles = {}


@dataclass
class ConversionConfig:
    """Configuration for markdown to Word conversion."""
    # Document settings
    font_name: str = "Calibri"
    font_size: int = 11
    line_spacing: float = 1.15
    
    # Heading settings
    heading_font: str = "Calibri"
    heading_colors: Dict[int, str] = None
    heading_sizes: Dict[int, int] = None
    heading_spacing_before: Dict[int, int] = None
    heading_spacing_after: Dict[int, int] = None
    
    # Paragraph settings
    paragraph_spacing_after: int = 6
    paragraph_first_line_indent: float = 0.0
    
    # List settings
    list_spacing: int = 0
    list_indent: float = 0.25
    
    # Table settings
    table_style: str = "Table Grid"
    table_autofit: bool = True
    
    # Code block settings
    code_font: str = "Consolas"
    code_size: int = 9
    code_background: str = "#F5F5F5"
    
    # Use Word built-in styles
    use_builtin_styles: bool = True
    custom_style_map: Dict[str, str] = None
    
    def __post_init__(self):
        if self.heading_colors is None:
            self.heading_colors = {
                1: "#2E75B6",  # Blue
                2: "#C55A11",  # Orange
                3: "#70AD47",  # Green
                4: "#7030A0",  # Purple
                5: "#264478",  # Dark Blue
                6: "#E7E6E6"   # Gray
            }
        
        if self.heading_sizes is None:
            self.heading_sizes = {
                1: 16, 2: 14, 3: 12, 4: 11, 5: 11, 6: 10
            }
        
        if self.heading_spacing_before is None:
            self.heading_spacing_before = {
                1: 12, 2: 10, 3: 8, 4: 6, 5: 6, 6: 4
            }
        
        if self.heading_spacing_after is None:
            self.heading_spacing_after = {
                1: 6, 2: 6, 3: 4, 4: 4, 5: 2, 6: 2
            }
        
        if self.custom_style_map is None:
            self.custom_style_map = {}


class FileSafetyManager:
    """Handles file safety operations: hashing, collision detection, backups."""
    
    def __init__(self, safety_config: SafetyConfig = None):
        self.config = safety_config or SafetyConfig()
    
    def calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA256 hash of file."""
        if not file_path.exists():
            return ""
        
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def detect_conversion_collision(self, source_file: Path, target_file: Path) -> bool:
        """Check if target file would create a conversion collision."""
        if not target_file.exists():
            return False
        
        # Check if target has same basename but different extension
        if source_file.stem == target_file.stem:
            source_hash = self.calculate_file_hash(source_file)
            target_hash = self.calculate_file_hash(target_file)
            
            # If hashes are different, it's a collision
            return source_hash != target_hash
        
        return False
    
    def create_backup(self, file_path: Path) -> Optional[Path]:
        """Create backup of existing file."""
        if not file_path.exists():
            return None
        
        backup_path = file_path.with_suffix(f"{file_path.suffix}{self.config.backup_suffix}")
        counter = 1
        
        while backup_path.exists():
            backup_path = file_path.with_suffix(f"{file_path.suffix}{self.config.backup_suffix}.{counter}")
            counter += 1
        
        try:
            shutil.copy2(file_path, backup_path)
            return backup_path
        except Exception as e:
            print(f"⚠️  Warning: Could not create backup: {e}")
            return None
    
    def confirm_overwrite(self, file_path: Path) -> bool:
        """Get user confirmation for file overwrite."""
        if not self.config.require_confirmation:
            return True
        
        response = input(f"⚠️  File '{file_path}' exists. Overwrite? [y/N]: ").lower().strip()
        return response in ['y', 'yes']
    
    def safe_write_check(self, source_file: Path, target_file: Path) -> Tuple[bool, str]:
        """
        Comprehensive safety check before writing.
        Returns (can_proceed, reason)
        """
        # Check for collision
        if self.detect_conversion_collision(source_file, target_file):
            return False, f"Collision detected: {target_file} exists with different content"
        
        # Check if target exists and get confirmation
        if target_file.exists():
            if self.config.prevent_overwrite:
                if not self.confirm_overwrite(target_file):
                    return False, "User cancelled overwrite"
            
            # Create backup if requested
            if self.config.create_backup:
                backup_path = self.create_backup(target_file)
                if backup_path:
                    print(f"✅ Backup created: {backup_path}")
        
        return True, "Safe to proceed"


class WordToMarkdownConverter:
    """Converts Word documents to Markdown with metadata preservation and safety."""
    
    def __init__(self, safety_manager: FileSafetyManager = None):
        self.metadata = FormatMetadata()
        self.current_line = 0
        self.safety = safety_manager or FileSafetyManager()
    
    def convert_docx_to_md(self, docx_path: str, md_path: str) -> bool:
        """Convert a Word document to Markdown with safety checks."""
        source_file = Path(docx_path)
        target_file = Path(md_path)
        
        if not source_file.exists():
            print(f"❌ Error: Source file {source_file} does not exist")
            return False
        
        # Safety check
        can_proceed, reason = self.safety.safe_write_check(source_file, target_file)
        if not can_proceed:
            print(f"❌ Safety check failed: {reason}")
            return False
        
        try:
            doc = Document(docx_path)
            markdown_content = self._extract_content_and_metadata(doc)
            
            # Add file hash to metadata
            self.metadata.file_hash = self.safety.calculate_file_hash(source_file)
            from datetime import datetime
            self.metadata.conversion_timestamp = datetime.now().isoformat()
            
            # Add metadata footer
            metadata_json = json.dumps(asdict(self.metadata), indent=2)
            full_content = f"{markdown_content}\n\n<!-- WORD_CONVERSION_METADATA\n{metadata_json}\n-->\n"
            
            with open(target_file, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            print(f"✅ Successfully converted {source_file} → {target_file}")
            print(f"📊 File hash: {self.metadata.file_hash[:16]}...")
            return True
            
        except Exception as e:
            print(f"❌ Error converting {source_file}: {e}")
            return False
    
    def _extract_content_and_metadata(self, doc: Document) -> str:
        """Extract content and build metadata."""
        content_lines = []
        self.current_line = 0
        
        for paragraph in doc.paragraphs:
            line_content = self._process_paragraph(paragraph)
            if line_content.strip():
                content_lines.append(line_content)
                self.current_line += 1
        
        # Process tables
        for table in doc.tables:
            table_md = self._process_table(table)
            if table_md:
                content_lines.append(table_md)
                self.current_line += table_md.count('\n')
        
        return '\n'.join(content_lines)
    
    def _process_paragraph(self, paragraph) -> str:
        """Process a paragraph and extract formatting."""
        if not paragraph.text.strip():
            return ""
        
        # Check for heading
        if paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('Title'):
            if paragraph.style.name == 'Title':
                level = 1
            else:
                level = int(re.findall(r'\d+', paragraph.style.name)[0]) if re.findall(r'\d+', paragraph.style.name) else 1
            
            self.metadata.heading_levels[self.current_line] = level
            return f"{'#' * level} {paragraph.text}"
        
        # Check for list items
        if any(list_style in paragraph.style.name for list_style in ['List', 'Bullet', 'Number']):
            list_info = {
                'line': self.current_line,
                'style': paragraph.style.name
                # Note: text content is already in markdown, no need to store again
            }
            self.metadata.lists.append(list_info)
            
            if any(bullet in paragraph.style.name for bullet in ['Bullet', 'bullet']):
                return f"- {paragraph.text}"
            else:
                return f"1. {paragraph.text}"
        
        # Process runs for inline formatting
        formatted_text = self._process_runs(paragraph.runs)
        
        # Check alignment and other paragraph properties
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            self.metadata.styles[str(self.current_line)] = "center"
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            self.metadata.styles[str(self.current_line)] = "right"
        
        return formatted_text
    
    def _process_runs(self, runs) -> str:
        """Process runs within a paragraph for inline formatting."""
        result = ""
        char_position = 0
        
        for run in runs:
            start_pos = char_position
            end_pos = char_position + len(run.text)
            
            text = run.text
            
            # Handle formatting - order matters for nested formatting
            formatting_applied = False
            
            # Handle superscript/subscript first
            if hasattr(run.font, 'superscript') and run.font.superscript:
                text = f"<sup>{text}</sup>"
                formatting_applied = True
            elif hasattr(run.font, 'subscript') and run.font.subscript:
                text = f"<sub>{text}</sub>"
                formatting_applied = True
            
            # Handle strikethrough
            if hasattr(run.font, 'strike') and run.font.strike:
                text = f"~~{text}~~"
                formatting_applied = True
            
            # Handle underline
            if run.underline:
                text = f"<u>{text}</u>"
                formatting_applied = True
            
            # Handle bold/italic combinations
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            # Handle hyperlinks
            if hasattr(run.element, 'xpath') and run.element.xpath('.//w:hyperlink'):
                hyperlink_info = {
                    'start': start_pos,
                    'end': end_pos,
                    'line': self.current_line
                    # Note: text content is preserved in markdown syntax
                }
                self.metadata.hyperlinks.append(hyperlink_info)
                text = f"[{run.text}](#{run.text.replace(' ', '-').lower()})"
            
            result += text
            char_position = end_pos
        
        return result
    
    def _process_table(self, table) -> str:
        """Convert Word table to Markdown table."""
        if not table.rows:
            return ""
        
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        
        if not table_data or not table_data[0]:
            return ""
        
        # Store table metadata
        table_info = {
            'line': self.current_line,
            'rows': len(table_data),
            'cols': len(table_data[0]),
            'data': table_data
        }
        self.metadata.tables.append(table_info)
        
        # Convert to Markdown table
        markdown_table = []
        markdown_table.append("| " + " | ".join(table_data[0]) + " |")
        markdown_table.append("| " + " | ".join(["---"] * len(table_data[0])) + " |")
        
        for row in table_data[1:]:
            markdown_table.append("| " + " | ".join(row) + " |")
        
        return "\n".join(markdown_table)


class MarkdownToWordConverter:
    """Converts Markdown to Word documents with configurable formatting and safety."""
    
    def __init__(self, config: ConversionConfig = None, template_path: str = None, safety_manager: FileSafetyManager = None):
        self.config = config or ConversionConfig()
        self.template_path = template_path
        self.safety = safety_manager or FileSafetyManager()
    
    def convert_md_to_docx(self, md_path: str, docx_path: str) -> bool:
        """Convert Markdown to Word document with safety checks."""
        source_file = Path(md_path)
        target_file = Path(docx_path)
        
        if not source_file.exists():
            print(f"❌ Error: Source file {source_file} does not exist")
            return False
        
        # Safety check
        can_proceed, reason = self.safety.safe_write_check(source_file, target_file)
        if not can_proceed:
            print(f"❌ Safety check failed: {reason}")
            return False
        
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Check for YAML frontmatter config
            frontmatter_config = self._extract_frontmatter_config(content)
            if frontmatter_config:
                self._update_config_from_dict(frontmatter_config)
                content = self._strip_frontmatter(content)
            
            # Extract metadata if present
            metadata = self._extract_metadata(content)
            markdown_content = self._strip_metadata(content)
            
            # Create Word document
            if self.template_path and Path(self.template_path).exists():
                doc = Document(self.template_path)
            else:
                doc = Document()
                self._setup_default_styles(doc)
            
            self._build_document(doc, markdown_content, metadata)
            
            doc.save(docx_path)
            
            # Calculate and display hash
            output_hash = self.safety.calculate_file_hash(target_file)
            print(f"✅ Successfully converted {source_file} → {target_file}")
            print(f"📊 Output hash: {output_hash[:16]}...")
            return True
            
        except Exception as e:
            print(f"❌ Error converting {source_file}: {e}")
            return False
    
    def _extract_frontmatter_config(self, content: str) -> Optional[Dict]:
        """Extract YAML frontmatter configuration."""
        if not yaml or not content.startswith('---\n'):
            return None
        
        try:
            end_marker = content.find('\n---\n', 4)
            if end_marker == -1:
                return None
            
            yaml_content = content[4:end_marker]
            return yaml.safe_load(yaml_content)
        except:
            return None
    
    def _strip_frontmatter(self, content: str) -> str:
        """Remove YAML frontmatter from content."""
        if not content.startswith('---\n'):
            return content
        
        end_marker = content.find('\n---\n', 4)
        if end_marker == -1:
            return content
        
        return content[end_marker + 5:]
    
    def _update_config_from_dict(self, config_dict: Dict) -> None:
        """Update configuration from dictionary."""
        for key, value in config_dict.items():
            if hasattr(self.config, key):
                setattr(self.config, key, value)
    
    def _extract_metadata(self, content: str) -> Optional[FormatMetadata]:
        """Extract metadata from Markdown content."""
        metadata_match = re.search(r'<!-- WORD_CONVERSION_METADATA\n(.*?)\n-->', content, re.DOTALL)
        if metadata_match:
            try:
                metadata_dict = json.loads(metadata_match.group(1))
                return FormatMetadata(**metadata_dict)
            except json.JSONDecodeError:
                pass
        return FormatMetadata()
    
    def _strip_metadata(self, content: str) -> str:
        """Remove metadata from content."""
        return re.sub(r'\n\n<!-- WORD_CONVERSION_METADATA.*?-->\n?$', '', content, flags=re.DOTALL)
    
    def _setup_default_styles(self, doc: Document) -> None:
        """Set up default document styles."""
        # Update Normal style
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.config.font_name
        normal_font.size = Pt(self.config.font_size)
        
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.space_after = Pt(self.config.paragraph_spacing_after)
        normal_paragraph.line_spacing = self.config.line_spacing
        normal_paragraph.first_line_indent = Inches(self.config.paragraph_first_line_indent)
        
        # Update heading styles
        for level in range(1, 7):
            try:
                heading_style = doc.styles[f'Heading {level}']
                self._configure_heading_style(heading_style, level)
            except KeyError:
                # Create heading style if it doesn't exist
                heading_style = doc.styles.add_style(f'Heading {level}', WD_STYLE_TYPE.PARAGRAPH)
                self._configure_heading_style(heading_style, level)
    
    def _configure_heading_style(self, style, level: int) -> None:
        """Configure a heading style."""
        font = style.font
        font.name = self.config.heading_font
        font.size = Pt(self.config.heading_sizes.get(level, 12))
        font.bold = True
        
        # Set color if specified
        if level in self.config.heading_colors:
            color_hex = self.config.heading_colors[level].replace('#', '')
            try:
                r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
                font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                pass
        
        # Set spacing
        paragraph = style.paragraph_format
        paragraph.space_before = Pt(self.config.heading_spacing_before.get(level, 6))
        paragraph.space_after = Pt(self.config.heading_spacing_after.get(level, 3))
        paragraph.keep_with_next = True
    
    def _build_document(self, doc: Document, content: str, metadata: FormatMetadata) -> None:
        """Build Word document from Markdown content and metadata."""
        # Parse markdown with extensions
        md = markdown.Markdown(extensions=['fenced_code', 'tables', 'toc'])
        
        lines = content.split('\n')
        current_list_type = None
        in_code_block = False
        code_block_content = []
        
        for line_num, line in enumerate(lines):
            line = line.rstrip()
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End of code block
                    self._add_code_block(doc, '\n'.join(code_block_content))
                    code_block_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_block_content.append(line)
                continue
            
            if not line.strip():
                if not in_code_block:
                    # Add empty paragraph for spacing
                    doc.add_paragraph()
                continue
            
            # Handle horizontal rules and dividers
            if self._is_horizontal_rule(line):
                self._add_horizontal_rule(doc)
                continue
            
            # Handle header boxes (equals dividers around text)
            if self._is_header_box_divider(line, lines, line_num):
                header_text = self._extract_header_box_text(lines, line_num)
                if header_text:
                    self._add_header_box(doc, header_text)
                    # Skip the header text line and closing divider
                    continue
            
            # Handle headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                if self.config.use_builtin_styles:
                    doc.add_heading(heading_text, level)
                else:
                    p = doc.add_paragraph(heading_text)
                    self._apply_custom_heading_format(p, level)
            
            # Handle lists
            elif line.strip().startswith(('-', '*', '+')):
                list_text = line.strip()[1:].strip()
                p = doc.add_paragraph(list_text, style='List Bullet')
                current_list_type = 'bullet'
            
            elif re.match(r'^\s*\d+\.', line):
                list_text = re.sub(r'^\s*\d+\.\s*', '', line)
                p = doc.add_paragraph(list_text, style='List Number')
                current_list_type = 'number'
            
            # Handle tables
            elif '|' in line and line.strip().startswith('|'):
                table_lines = [line]
                # Collect all table lines
                for next_line_num in range(line_num + 1, len(lines)):
                    next_line = lines[next_line_num]
                    if '|' in next_line and next_line.strip().startswith('|'):
                        table_lines.append(next_line)
                    else:
                        break
                
                if len(table_lines) >= 2:  # Header + separator minimum
                    self._add_markdown_table(doc, table_lines)
                    # Skip the lines we just processed
                    line_num += len(table_lines) - 1
            
            # Handle blockquotes
            elif line.strip().startswith('>'):
                quote_text = line.strip()[1:].strip()
                p = doc.add_paragraph(quote_text)
                p.style = 'Quote'
            
            # Regular paragraphs
            else:
                p = doc.add_paragraph()
                self._apply_inline_formatting(p, line)
                current_list_type = None
    
    def _add_code_block(self, doc: Document, code_content: str) -> None:
        """Add a code block to the document."""
        p = doc.add_paragraph(code_content)
        
        # Style the code block
        font = p.runs[0].font if p.runs else p.style.font
        font.name = self.config.code_font
        font.size = Pt(self.config.code_size)
        
        # Add background color (limited support in python-docx)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    def _apply_custom_heading_format(self, paragraph, level: int) -> None:
        """Apply custom heading formatting."""
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        
        font = run.font
        font.name = self.config.heading_font
        font.size = Pt(self.config.heading_sizes.get(level, 12))
        font.bold = True
        
        if level in self.config.heading_colors:
            color_hex = self.config.heading_colors[level].replace('#', '')
            try:
                r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
                font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                pass
        
        paragraph.paragraph_format.space_before = Pt(self.config.heading_spacing_before.get(level, 6))
        paragraph.paragraph_format.space_after = Pt(self.config.heading_spacing_after.get(level, 3))
    
    def _apply_inline_formatting(self, paragraph, text: str) -> None:
        """Apply inline formatting to paragraph text."""
        # Parse inline markdown
        current_pos = 0
        
        # Combined pattern for bold, italic, code
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))'
        
        for match in re.finditer(pattern, text):
            # Add text before formatting
            if match.start() > current_pos:
                paragraph.add_run(text[current_pos:match.start()])
            
            # Determine formatting type and add formatted run
            full_match = match.group(1)
            if full_match.startswith('***') and full_match.endswith('***'):
                # Bold + italic
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif full_match.startswith('**') and full_match.endswith('**'):
                # Bold
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif full_match.startswith('*') and full_match.endswith('*'):
                # Italic
                run = paragraph.add_run(match.group(4))
                run.italic = True
            elif full_match.startswith('`') and full_match.endswith('`'):
                # Inline code
                run = paragraph.add_run(match.group(5))
                run.font.name = self.config.code_font
                run.font.size = Pt(self.config.code_size)
            elif '[' in full_match and '](' in full_match:
                # Hyperlink
                run = paragraph.add_run(match.group(6))
                # Note: Creating actual hyperlinks in python-docx is complex
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            
            current_pos = match.end()
        
        # Add remaining text
        if current_pos < len(text):
            paragraph.add_run(text[current_pos:])
    
    def _add_markdown_table(self, doc: Document, table_lines: List[str]) -> None:
        """Add a markdown table to the document."""
        # Parse table data
        rows = []
        for line in table_lines:
            if '---' in line:  # Skip separator line
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        # Create Word table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = self.config.table_style
        
        if self.config.table_autofit:
            table.autofit = True
        
        # Populate table
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    # Apply inline formatting to cell text
                    cell.text = cell_data
                    if row_idx == 0:  # Header row
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
    
    def _is_horizontal_rule(self, line: str) -> bool:
        """Check if line is a horizontal divider pattern."""
        stripped = line.strip()
        if not stripped:
            return False
        
        # Standard markdown horizontal rules
        if stripped in ['---', '***', '___']:
            return True
        
        # Long divider lines (10+ characters of same symbol)
        if len(stripped) >= 10:
            # Check for repeated dashes, equals, or unicode box drawing
            if all(c == '-' for c in stripped):  # ----------------
                return True
            if all(c == '=' for c in stripped):  # ================
                return True
            if all(c == '─' for c in stripped):  # ────────────────
                return True
        
        return False
    
    def _add_horizontal_rule(self, doc: Document) -> None:
        """Add a horizontal rule to the Word document."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create paragraph with underline to simulate horizontal line
        p = doc.add_paragraph("_" * 50)  # Create underline characters
        
        # Style as a horizontal line with minimal spacing
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)  # Reduced from 12
        p.paragraph_format.space_before = Pt(3)  # Reduced from 6
        
        # Style the text to look like a line
        if p.runs:
            run = p.runs[0]
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
            run.font.size = Pt(8)
    
    def _is_header_box_divider(self, line: str, lines: List[str], line_num: int) -> bool:
        """Check if line is start of a header box pattern (equals dividers)."""
        stripped = line.strip()
        
        # Must be long line of equals signs
        if len(stripped) >= 20 and all(c == '=' for c in stripped):
            # Check if there's a text line and closing divider following
            if line_num + 2 < len(lines):
                text_line = lines[line_num + 1].strip()
                closing_line = lines[line_num + 2].strip()
                
                # Text line should not be empty and should not be another divider
                if text_line and not all(c in '=-_' for c in text_line if c.strip()):
                    # Closing line should also be equals divider
                    if len(closing_line) >= 20 and all(c == '=' for c in closing_line):
                        return True
        
        return False
    
    def _extract_header_box_text(self, lines: List[str], line_num: int) -> Optional[str]:
        """Extract text from header box pattern and mark lines as processed."""
        if line_num + 2 < len(lines):
            header_text = lines[line_num + 1].strip()
            
            # Mark the header text line and closing divider as processed
            # by setting them to empty (they'll be skipped in the main loop)
            lines[line_num + 1] = ""
            lines[line_num + 2] = ""
            
            return header_text
        return None
    
    def _add_header_box(self, doc: Document, header_text: str) -> None:
        """Add a bordered header box to the Word document."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create the header paragraph with minimal spacing
        p = doc.add_paragraph(header_text)
        
        # Style the header box
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add top and bottom borders
        p.paragraph_format.border_top.color.rgb = RGBColor(0, 0, 0)  # Black
        p.paragraph_format.border_top.width = Pt(2)
        p.paragraph_format.border_bottom.color.rgb = RGBColor(0, 0, 0)  # Black  
        p.paragraph_format.border_bottom.width = Pt(2)
        
        # Minimal padding
        p.paragraph_format.space_before = Pt(4)  # Reduced from 12
        p.paragraph_format.space_after = Pt(4)   # Reduced from 12
        
        # Style the text
        if p.runs:
            run = p.runs[0]
            run.font.bold = True
            run.font.size = Pt(14)
        else:
            # If no runs, add the text as a run
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(14)


def load_config_file(config_path: str) -> ConversionConfig:
    """Load configuration from file."""
    if not Path(config_path).exists():
        print(f"⚠️  Config file {config_path} not found, using defaults")
        return ConversionConfig()
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            if config_path.endswith('.json'):
                config_dict = json.load(f)
            elif config_path.endswith(('.yml', '.yaml')) and yaml:
                config_dict = yaml.safe_load(f)
            else:
                print("⚠️  Unsupported config format, using defaults")
                return ConversionConfig()
        
        # Create config object
        config = ConversionConfig()
        for key, value in config_dict.items():
            if hasattr(config, key):
                setattr(config, key, value)
        
        return config
    except Exception as e:
        print(f"⚠️  Error loading config: {e}, using defaults")
        return ConversionConfig()


def create_sample_config(config_path: str) -> None:
    """Create a sample configuration file."""
    config = ConversionConfig()
    config_dict = asdict(config)
    
    try:
        if config_path.endswith('.json'):
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config_dict, f, indent=2)
        elif config_path.endswith(('.yml', '.yaml')) and yaml:
            with open(config_path, 'w', encoding='utf-8') as f:
                yaml.dump(config_dict, f, default_flow_style=False)
        else:
            print("❌ Unsupported config format")
            return
        
        print(f"✅ Sample configuration created at {config_path}")
    except Exception as e:
        print(f"❌ Error creating config: {e}")


def main():
    """Main function to handle command line arguments with safety features."""
    parser = argparse.ArgumentParser(
        description='Word - Safe document conversion tool with hash validation',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  word document.docx document.md          # Convert Word to Markdown
  word document.md document.docx          # Convert Markdown to Word
  word --force document.md document.docx  # Skip confirmation prompts
  word --no-backup document.md output.docx # Skip backup creation
  word --create-config config.yaml        # Create sample configuration

Safety Features:
  - Hash validation prevents data loss
  - Automatic backup creation
  - Collision detection
  - Confirmation prompts for overwrites
        """
    )
    
    parser.add_argument('input_file', nargs='?', help='Input file path')
    parser.add_argument('output_file', nargs='?', help='Output file path')
    parser.add_argument('--direction', choices=['docx2md', 'md2docx'], 
                       help='Conversion direction (auto-detected if not specified)')
    parser.add_argument('--config', help='Configuration file path (JSON or YAML)')
    parser.add_argument('--template', help='Word template file for MD to DOCX conversion')
    parser.add_argument('--create-config', help='Create sample configuration file')
    
    # Safety options
    parser.add_argument('--force', action='store_true', help='Skip confirmation prompts')
    parser.add_argument('--no-backup', action='store_true', help='Skip backup creation')
    parser.add_argument('--no-hash-check', action='store_true', help='Skip hash validation')
    
    args = parser.parse_args()
    
    # Handle config creation
    if args.create_config:
        create_sample_config(args.create_config)
        return
    
    if not args.input_file or not args.output_file:
        parser.print_help()
        sys.exit(1)
    
    input_path = Path(args.input_file)
    output_path = Path(args.output_file)
    
    if not input_path.exists():
        print(f"❌ Error: Input file {input_path} does not exist")
        sys.exit(1)
    
    # Configure safety settings
    safety_config = SafetyConfig(
        require_confirmation=not args.force,
        create_backup=not args.no_backup,
        check_hash=not args.no_hash_check,
        prevent_overwrite=True
    )
    
    safety_manager = FileSafetyManager(safety_config)
    
    # Load configuration
    config = load_config_file(args.config) if args.config else ConversionConfig()
    
    # Auto-detect conversion direction
    if args.direction:
        direction = args.direction
    else:
        if input_path.suffix.lower() == '.docx':
            direction = 'docx2md'
        elif input_path.suffix.lower() == '.md':
            direction = 'md2docx'
        else:
            print("❌ Error: Cannot auto-detect conversion direction. Please specify --direction")
            sys.exit(1)
    
    try:
        if direction == 'docx2md':
            converter = WordToMarkdownConverter(safety_manager)
            success = converter.convert_docx_to_md(str(input_path), str(output_path))
        else:
            converter = MarkdownToWordConverter(config, args.template, safety_manager)
            success = converter.convert_md_to_docx(str(input_path), str(output_path))
        
        sys.exit(0 if success else 1)
            
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()